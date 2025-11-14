from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file, jsonify
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
import os
from datetime import datetime
import pandas as pd
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as ReportLabImage
from reportlab.lib.styles import getSampleStyleSheet
import io
import shutil
import zipfile
import PyPDF2
import openpyxl
from PIL import Image as PILImage
import docx2txt
import base64

# Import docx in a way that Pylance accepts
try:
    import docx
    from docx.api import Document  # type: ignore
except ImportError:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        Document = None
        print("Warning: python-docx not available. Word export will not work.")

app = Flask(__name__)
app.config['SECRET_KEY'] = 'your-secret-key-here-change-in-production'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///database.db'
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

db = SQLAlchemy(app)

# Allowed file extensions
ALLOWED_EXTENSIONS = {'pdf', 'doc', 'docx', 'xls', 'xlsx', 'jpg', 'jpeg', 'png'}

class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password = db.Column(db.String(120), nullable=False)
    role = db.Column(db.String(20), nullable=False)
    documentations = db.relationship('Documentation', backref='author', lazy=True)

class Documentation(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(200), nullable=False)
    content = db.Column(db.Text)
    file_path = db.Column(db.String(300))
    doc_type = db.Column(db.String(50), nullable=False)
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)

def allowed_file(filename):
    if '.' not in filename:
        return False
    return filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def can_edit_delete(doc):
    """Check if current user can edit/delete the document - ONLY ADMIN"""
    if 'user_id' not in session:
        return False
    
    # ONLY Admin can edit/delete documents
    return session.get('role') == 'admin'

def can_view(doc):
    """Check if current user can view the document"""
    if 'user_id' not in session:
        return False
    
    # Both admin and users can view all documents
    return True

def get_image_base64(file_path):
    """Convert image to base64 for display in HTML"""
    try:
        with open(file_path, 'rb') as img_file:
            return base64.b64encode(img_file.read()).decode('utf-8')
    except Exception as e:
        return None

def pdf_to_excel(pdf_path, title):
    """Convert PDF to Excel while preserving structure"""
    try:
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "PDF Content"
        
        # Add title
        sheet['A1'] = title
        sheet['A1'].font = openpyxl.styles.Font(size=14, bold=True)
        
        # Extract text from PDF with structure
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            row_num = 3
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                
                if text.strip():
                    # Add page header
                    sheet[f'A{row_num}'] = f"Page {page_num + 1}"
                    sheet[f'A{row_num}'].font = openpyxl.styles.Font(bold=True)
                    row_num += 1
                    
                    # Add content lines
                    lines = text.split('\n')
                    for line in lines:
                        if line.strip():
                            sheet[f'A{row_num}'] = line.strip()
                            row_num += 1
                    
                    row_num += 1  # Add space between pages
        
        # Auto-adjust column width
        sheet.column_dimensions['A'].width = 50
        
        output = io.BytesIO()
        workbook.save(output)
        output.seek(0)
        return output
        
    except Exception as e:
        # Fallback: create simple Excel with error message
        return create_fallback_excel(title, f"Error processing PDF: {str(e)}")

def pdf_to_word(pdf_path, title):
    """Convert PDF to Word while preserving structure"""
    try:
        if Document is None:
            raise Exception("python-docx not available")
        
        doc = Document()
        doc.add_heading(title, 0)
        
        # Extract text from PDF
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                
                if text.strip():
                    doc.add_heading(f"Page {page_num + 1}", level=1)
                    lines = text.split('\n')
                    for line in lines:
                        if line.strip():
                            doc.add_paragraph(line.strip())
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        # Fallback: create simple Word with error message
        return create_fallback_word(title, f"Error processing PDF: {str(e)}")

def image_to_pdf(image_path, title):
    """Convert image to PDF"""
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Add title
        title_para = Paragraph(title, styles['Title'])
        story.append(title_para)
        story.append(Spacer(1, 12))
        
        # Add image
        img = ReportLabImage(image_path, width=400, height=300)
        story.append(img)
        
        doc.build(story)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        return create_fallback_pdf(title, f"Error processing image: {str(e)}")

def image_to_excel(image_path, title):
    """Convert image to Excel with image info"""
    try:
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "Image Info"
        
        # Add title
        sheet['A1'] = title
        sheet['A1'].font = openpyxl.styles.Font(size=14, bold=True)
        
        # Add image information
        with PILImage.open(image_path) as img:
            sheet['A3'] = "Image Information:"
            sheet['A3'].font = openpyxl.styles.Font(bold=True)
            sheet['A4'] = f"Format: {img.format}"
            sheet['A5'] = f"Size: {img.size[0]} x {img.size[1]} pixels"
            sheet['A6'] = f"Mode: {img.mode}"
            sheet['A7'] = f"File: {os.path.basename(image_path)}"
        
        # Auto-adjust column width
        sheet.column_dimensions['A'].width = 30
        
        output = io.BytesIO()
        workbook.save(output)
        output.seek(0)
        return output
        
    except Exception as e:
        return create_fallback_excel(title, f"Error processing image: {str(e)}")

def image_to_word(image_path, title):
    """Convert image to Word with image"""
    try:
        if Document is None:
            raise Exception("python-docx not available")
        
        doc = Document()
        doc.add_heading(title, 0)
        
        # Add image information
        with PILImage.open(image_path) as img:
            doc.add_paragraph(f"Image Information:")
            doc.add_paragraph(f"Format: {img.format}")
            doc.add_paragraph(f"Size: {img.size[0]} x {img.size[1]} pixels")
            doc.add_paragraph(f"Mode: {img.mode}")
            doc.add_paragraph(f"File: {os.path.basename(image_path)}")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        return create_fallback_word(title, f"Error processing image: {str(e)}")

def excel_to_pdf(excel_path, title):
    """Convert Excel to PDF"""
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Add title
        title_para = Paragraph(title, styles['Title'])
        story.append(title_para)
        story.append(Spacer(1, 12))
        
        # Read Excel data
        workbook = openpyxl.load_workbook(excel_path)
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            
            # Add sheet name
            story.append(Paragraph(f"Sheet: {sheet_name}", styles['Heading2']))
            story.append(Spacer(1, 6))
            
            # Add data (limited to first 50 rows for PDF)
            row_count = 0
            for row in sheet.iter_rows(values_only=True):
                if any(cell is not None for cell in row):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    story.append(Paragraph(row_text, styles['Normal']))
                    row_count += 1
                
                if row_count >= 50:  # Limit rows in PDF
                    story.append(Paragraph("... (more rows in original file)", styles['Italic']))
                    break
            
            story.append(Spacer(1, 12))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        return create_fallback_pdf(title, f"Error processing Excel: {str(e)}")

def excel_to_word(excel_path, title):
    """Convert Excel to Word"""
    try:
        if Document is None:
            raise Exception("python-docx not available")
        
        doc = Document()
        doc.add_heading(title, 0)
        
        # Read Excel data
        workbook = openpyxl.load_workbook(excel_path)
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            
            doc.add_heading(f"Sheet: {sheet_name}", level=1)
            
            # Add data (limited to first 30 rows for Word)
            row_count = 0
            for row in sheet.iter_rows(values_only=True):
                if any(cell is not None for cell in row):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    doc.add_paragraph(row_text)
                    row_count += 1
                
                if row_count >= 30:  # Limit rows in Word
                    doc.add_paragraph("... (more rows in original file)")
                    break
            
            doc.add_paragraph()  # Empty line between sheets
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        return create_fallback_word(title, f"Error processing Excel: {str(e)}")

def word_to_pdf(word_path, title):
    """Convert Word to PDF"""
    try:
        # Extract text from Word
        text = docx2txt.process(word_path)
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Add title
        title_para = Paragraph(title, styles['Title'])
        story.append(title_para)
        story.append(Spacer(1, 12))
        
        # Add content
        if text.strip():
            lines = text.split('\n')
            for line in lines:
                if line.strip():
                    story.append(Paragraph(line.strip(), styles['Normal']))
                    story.append(Spacer(1, 6))
        else:
            story.append(Paragraph("No content found in document", styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        return create_fallback_pdf(title, f"Error processing Word document: {str(e)}")

def word_to_excel(word_path, title):
    """Convert Word to Excel"""
    try:
        # Extract text from Word
        text = docx2txt.process(word_path)
        
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "Document Content"
        
        # Add title
        sheet['A1'] = title
        sheet['A1'].font = openpyxl.styles.Font(size=14, bold=True)
        
        # Add content
        if text.strip():
            lines = text.split('\n')
            row_num = 3
            for i, line in enumerate(lines):
                if line.strip():
                    sheet[f'A{row_num}'] = line.strip()
                    row_num += 1
        else:
            sheet['A3'] = "No content found in document"
        
        # Auto-adjust column width
        sheet.column_dimensions['A'].width = 50
        
        output = io.BytesIO()
        workbook.save(output)
        output.seek(0)
        return output
        
    except Exception as e:
        return create_fallback_excel(title, f"Error processing Word document: {str(e)}")

def manual_to_pdf(content, title):
    """Convert manual content to PDF"""
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Add title
        title_para = Paragraph(title, styles['Title'])
        story.append(title_para)
        story.append(Spacer(1, 12))
        
        # Add content
        if content and content.strip():
            lines = content.split('\n')
            for line in lines:
                if line.strip():
                    story.append(Paragraph(line.strip(), styles['Normal']))
                    story.append(Spacer(1, 6))
        else:
            story.append(Paragraph("No content available", styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        return create_fallback_pdf(title, f"Error processing manual content: {str(e)}")

def manual_to_excel(content, title):
    """Convert manual content to Excel"""
    try:
        workbook = openpyxl.Workbook()
        sheet = workbook.active
        sheet.title = "Document Content"
        
        # Add title
        sheet['A1'] = title
        sheet['A1'].font = openpyxl.styles.Font(size=14, bold=True)
        
        # Add content
        if content and content.strip():
            lines = content.split('\n')
            row_num = 3
            for i, line in enumerate(lines):
                if line.strip():
                    sheet[f'A{row_num}'] = line.strip()
                    row_num += 1
        else:
            sheet['A3'] = "No content available"
        
        # Auto-adjust column width
        sheet.column_dimensions['A'].width = 50
        
        output = io.BytesIO()
        workbook.save(output)
        output.seek(0)
        return output
        
    except Exception as e:
        return create_fallback_excel(title, f"Error processing manual content: {str(e)}")

def manual_to_word(content, title):
    """Convert manual content to Word"""
    try:
        if Document is None:
            raise Exception("python-docx not available")
        
        doc = Document()
        doc.add_heading(title, 0)
        
        # Add content
        if content and content.strip():
            lines = content.split('\n')
            for line in lines:
                if line.strip():
                    doc.add_paragraph(line.strip())
        else:
            doc.add_paragraph("No content available")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        return create_fallback_word(title, f"Error processing manual content: {str(e)}")

def create_fallback_pdf(title, message):
    """Create a fallback PDF when conversion fails"""
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)
    p.setFont("Helvetica-Bold", 16)
    p.drawString(100, 750, title)
    p.setFont("Helvetica", 12)
    p.drawString(100, 700, message)
    p.save()
    buffer.seek(0)
    return buffer

def create_fallback_excel(title, message):
    """Create a fallback Excel when conversion fails"""
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet['A1'] = title
    sheet['A1'].font = openpyxl.styles.Font(size=14, bold=True)
    sheet['A3'] = message
    output = io.BytesIO()
    workbook.save(output)
    output.seek(0)
    return output

def create_fallback_word(title, message):
    """Create a fallback Word when conversion fails"""
    if Document is None:
        return create_fallback_excel(title, message + " (Word export not available)")
    
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(message)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@app.route('/')
def index():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    return redirect(url_for('dashboard'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        user = User.query.filter_by(username=username).first()
        
        if user and check_password_hash(user.password, password):
            session['user_id'] = user.id
            session['username'] = user.username
            session['role'] = user.role
            flash('Login successful!', 'success')
            return redirect(url_for('dashboard'))
        else:
            flash('Invalid username or password', 'error')
    
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    flash('You have been logged out', 'info')
    return redirect(url_for('login'))

@app.route('/dashboard')
def dashboard():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    # Both admin and users can see ALL documents
    docs = Documentation.query.all()
    
    # Prepare documents with additional data for display
    docs_with_content = []
    for doc in docs:
        doc_data = {
            'doc': doc,
            'image_preview': None
        }
        
        # If it's an image, get base64 for preview
        if doc.doc_type == 'image' and doc.file_path and os.path.exists(doc.file_path):
            doc_data['image_preview'] = get_image_base64(doc.file_path)
        
        docs_with_content.append(doc_data)
    
    return render_template('dashboard.html', docs=docs_with_content)

@app.route('/add_doc', methods=['GET', 'POST'])
def add_doc():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        title = request.form['title']
        doc_type = request.form['doc_type']
        content = request.form.get('content', '')
        file = request.files.get('file')
        
        file_path = None
        if file and file.filename:
            if allowed_file(file.filename):
                filename = secure_filename(file.filename)
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(file_path)
            else:
                flash('File type not allowed. Allowed types: PDF, Word, Excel, Images', 'error')
                return redirect(url_for('add_doc'))
        
        new_doc = Documentation(
            title=title,
            content=content,
            file_path=file_path,
            doc_type=doc_type,
            user_id=session['user_id']
        )
        
        db.session.add(new_doc)
        db.session.commit()
        flash('Documentation added successfully!', 'success')
        return redirect(url_for('dashboard'))
    
    return render_template('add_doc.html')

@app.route('/edit_doc/<int:doc_id>', methods=['GET', 'POST'])
def edit_doc(doc_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    doc = Documentation.query.get_or_404(doc_id)
    
    # Check permission - ONLY admin can edit
    if not can_edit_delete(doc):
        flash('You do not have permission to edit documents', 'error')
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        doc.title = request.form['title']
        doc.doc_type = request.form['doc_type']
        doc.content = request.form.get('content', '')
        
        # Handle file update
        file = request.files.get('file')
        if file and file.filename:
            if allowed_file(file.filename):
                # Delete old file if exists
                if doc.file_path and os.path.exists(doc.file_path):
                    os.remove(doc.file_path)
                
                # Save new file
                filename = secure_filename(file.filename)
                doc.file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(doc.file_path)
            else:
                flash('File type not allowed. Allowed types: PDF, Word, Excel, Images', 'error')
                return redirect(url_for('edit_doc', doc_id=doc_id))
        
        db.session.commit()
        flash('Document updated successfully!', 'success')
        return redirect(url_for('dashboard'))
    
    return render_template('edit_doc.html', doc=doc)

@app.route('/delete_doc/<int:doc_id>', methods=['POST'])
def delete_doc(doc_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    doc = Documentation.query.get_or_404(doc_id)
    
    # Check permission - ONLY admin can delete
    if not can_edit_delete(doc):
        flash('You do not have permission to delete documents', 'error')
        return redirect(url_for('dashboard'))
    
    # Delete associated file
    if doc.file_path and os.path.exists(doc.file_path):
        os.remove(doc.file_path)
    
    db.session.delete(doc)
    db.session.commit()
    flash('Document deleted successfully!', 'success')
    return redirect(url_for('dashboard'))

@app.route('/view_doc/<int:doc_id>')
def view_doc(doc_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    doc = Documentation.query.get_or_404(doc_id)
    
    # Both admin and users can view all documents
    if not can_view(doc):
        flash('You do not have permission to view this document', 'error')
        return redirect(url_for('dashboard'))
    
    # Prepare content for display
    image_preview = None
    
    if doc.file_path and os.path.exists(doc.file_path) and doc.doc_type == 'image':
        image_preview = get_image_base64(doc.file_path)
    
    return render_template('view_doc.html', 
                         doc=doc, 
                         image_preview=image_preview)

@app.route('/download_file/<int:doc_id>')
def download_file(doc_id):
    """Download individual uploaded file in its original format"""
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    doc = Documentation.query.get_or_404(doc_id)
    
    # Both admin and users can download all files (since they can view all)
    if not can_view(doc):
        flash('You do not have permission to download this file', 'error')
        return redirect(url_for('dashboard'))
    
    if not doc.file_path or not os.path.exists(doc.file_path):
        flash('File not found', 'error')
        return redirect(url_for('dashboard'))
    
    return send_file(
        doc.file_path,
        as_attachment=True,
        download_name=os.path.basename(doc.file_path)
    )

@app.route('/convert_doc/<int:doc_id>/<target_format>')
def convert_doc(doc_id, target_format):
    """Convert document to different format and download"""
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    doc = Documentation.query.get_or_404(doc_id)
    
    if not can_view(doc):
        flash('You do not have permission to convert this document', 'error')
        return redirect(url_for('dashboard'))
    
    try:
        original_type = doc.doc_type
        buffer = None
        
        # Handle manual input documents (no file_path)
        if original_type == 'manual':
            content = doc.content or "No content available"
            if target_format == 'pdf':
                buffer = manual_to_pdf(content, doc.title)
            elif target_format == 'excel':
                buffer = manual_to_excel(content, doc.title)
            elif target_format == 'word':
                buffer = manual_to_word(content, doc.title)
            else:
                flash(f'Conversion from manual input to {target_format} is not supported', 'error')
                return redirect(url_for('view_doc', doc_id=doc_id))
        
        # Handle file-based documents
        else:
            if not doc.file_path or not os.path.exists(doc.file_path):
                flash('Original file not found', 'error')
                return redirect(url_for('view_doc', doc_id=doc_id))
            
            if original_type == 'pdf':
                if target_format == 'excel':
                    buffer = pdf_to_excel(doc.file_path, doc.title)
                elif target_format == 'word':
                    buffer = pdf_to_word(doc.file_path, doc.title)
                elif target_format == 'pdf':
                    # Same format, just download original
                    return send_file(
                        doc.file_path,
                        as_attachment=True,
                        download_name=f"{doc.title}.pdf"
                    )
            
            elif original_type == 'image':
                if target_format == 'pdf':
                    buffer = image_to_pdf(doc.file_path, doc.title)
                elif target_format == 'excel':
                    buffer = image_to_excel(doc.file_path, doc.title)
                elif target_format == 'word':
                    buffer = image_to_word(doc.file_path, doc.title)
            
            elif original_type == 'excel':
                if target_format == 'pdf':
                    buffer = excel_to_pdf(doc.file_path, doc.title)
                elif target_format == 'word':
                    buffer = excel_to_word(doc.file_path, doc.title)
                elif target_format == 'excel':
                    return send_file(
                        doc.file_path,
                        as_attachment=True,
                        download_name=f"{doc.title}.xlsx"
                    )
            
            elif original_type == 'word':
                if target_format == 'pdf':
                    buffer = word_to_pdf(doc.file_path, doc.title)
                elif target_format == 'excel':
                    buffer = word_to_excel(doc.file_path, doc.title)
                elif target_format == 'word':
                    return send_file(
                        doc.file_path,
                        as_attachment=True,
                        download_name=f"{doc.title}.docx"
                    )
        
        if buffer is None:
            flash(f'Conversion from {original_type} to {target_format} is not supported', 'error')
            return redirect(url_for('view_doc', doc_id=doc_id))
        
        # Set appropriate filename and mimetype
        if target_format == 'pdf':
            filename = f"{doc.title}.pdf"
            mimetype = 'application/pdf'
        elif target_format == 'excel':
            filename = f"{doc.title}.xlsx"
            mimetype = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        elif target_format == 'word':
            filename = f"{doc.title}.docx"
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            filename = f"{doc.title}.{target_format}"
            mimetype = 'application/octet-stream'
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=filename,
            mimetype=mimetype
        )
        
    except Exception as e:
        flash(f'Error converting document: {str(e)}', 'error')
        return redirect(url_for('view_doc', doc_id=doc_id))

def init_db():
    with app.app_context():
        db.create_all()
        if not User.query.filter_by(username='admin').first():
            admin_user = User(
                username='admin',
                password=generate_password_hash('admin123'),
                role='admin'
            )
            db.session.add(admin_user)
        
        if not User.query.filter_by(username='user').first():
            regular_user = User(
                username='user',
                password=generate_password_hash('user123'),
                role='user'
            )
            db.session.add(regular_user)
        
        db.session.commit()

if __name__ == '__main__':
    if not os.path.exists('uploads'):
        os.makedirs('uploads')
    init_db()
    app.run(debug=True)